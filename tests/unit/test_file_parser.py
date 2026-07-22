"""Unit tests for user-uploaded document parsing."""

from __future__ import annotations

from io import BytesIO
from typing import ClassVar

import pytest
from docx import Document
from pypdf import PdfWriter
from utils import file_parser


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("一级标题", level=1)
    document.add_heading("三级标题", level=3)
    document.add_paragraph("正文内容用于验证 Word 文档解析。")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_parse_utf8_markdown_and_gbk_text() -> None:
    assert file_parser.parse("article.MD", "# 标题\n\n正文".encode()) == "# 标题\n\n正文"
    assert file_parser.parse("article.txt", "中文 GBK 文本".encode("gbk")) == "中文 GBK 文本"


def test_parse_pdf_joins_non_empty_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    class Page:
        def __init__(self, text: str | None):
            self.text = text

        def extract_text(self) -> str | None:
            return self.text

    class Reader:
        is_encrypted = False
        pages: ClassVar[list[Page]] = [Page("第一页"), Page(None), Page(" 第二页 ")]

    monkeypatch.setattr(file_parser, "PdfReader", lambda _stream: Reader())

    assert file_parser.parse("report.pdf", b"fake-pdf") == "第一页\n\n第二页"


def test_encrypted_pdf_raises_clear_error() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    stream = BytesIO()
    writer.write(stream)

    with pytest.raises(file_parser.FileParseError, match="加密 PDF"):
        file_parser.parse("secret.pdf", stream.getvalue())


def test_parse_docx_converts_heading_styles_to_markdown() -> None:
    parsed = file_parser.parse("report.docx", _docx_bytes())

    assert parsed == "# 一级标题\n\n### 三级标题\n\n正文内容用于验证 Word 文档解析。"


def test_corrupt_docx_and_unsupported_extension_raise_parse_error() -> None:
    with pytest.raises(file_parser.FileParseError, match="docx 解析失败"):
        file_parser.parse("broken.docx", b"not-a-docx")
    with pytest.raises(file_parser.FileParseError, match="不支持的文件类型"):
        file_parser.parse("malware.exe", b"content")
