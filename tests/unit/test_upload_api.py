"""API tests for authenticated article uploads."""

from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock

import pytest
from api import upload as upload_api
from auth.deps import AuthError, auth_error_handler, get_current_user
from docx import Document
from fastapi import FastAPI
from httpx import ASGITransport, AsyncClient
from utils.file_parser import MAX_FILE_SIZE


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("上传文档标题", level=1)
    document.add_paragraph("这是用于上传接口测试的正文内容。" * 5)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _make_app(*, authenticated: bool = True, duplicate: bool = False):
    articles = MagicMock()
    articles.find_one = AsyncMock(return_value={"url_hash": "existing"} if duplicate else None)
    articles.insert_one = AsyncMock(return_value=SimpleNamespace(inserted_id="article-id"))
    activities = MagicMock()
    activities.insert_one = AsyncMock(return_value=SimpleNamespace(inserted_id="activity-id"))
    database = MagicMock()
    database.__getitem__.side_effect = lambda name: {
        "articles": articles,
        "user_activities": activities,
    }[name]

    app = FastAPI()
    app.state.db = database
    app.add_exception_handler(AuthError, auth_error_handler)
    app.include_router(upload_api.router)
    if authenticated:

        async def current_user() -> str:
            return "user-10"

        app.dependency_overrides[get_current_user] = current_user
    return app, articles, activities


async def _post(app: FastAPI, filename: str, content: bytes, *, title: str | None = None):
    data = {"title": title} if title is not None else None
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        return await client.post(
            "/api/upload/article",
            data=data,
            files={"file": (filename, content, "application/octet-stream")},
        )


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("filename", "content"),
    [
        ("article.txt", ("普通文本上传内容。" * 10).encode()),
        ("article.md", ("# Markdown 标题\n\n正文内容。" * 6).encode()),
        ("article.docx", _docx_bytes()),
    ],
    ids=("txt", "markdown", "docx"),
)
async def test_upload_supported_text_and_docx_files(filename: str, content: bytes) -> None:
    app, articles, activities = _make_app()

    response = await _post(app, filename, content)

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["source_type"] == "user_upload"
    assert payload["content_length"] >= 50
    document = articles.insert_one.await_args.args[0]
    assert document["uploaded_by"] == "user-10"
    assert document["original_filename"] == filename
    assert document["pipeline_status"] == "pending"
    assert document["content_md"]
    activities.insert_one.assert_awaited_once()


@pytest.mark.asyncio
async def test_upload_pdf_uses_parser_and_custom_title(monkeypatch: pytest.MonkeyPatch) -> None:
    app, articles, _activities = _make_app()
    monkeypatch.setattr(upload_api, "parse", lambda _name, _content: "PDF 提取正文。" * 10)

    response = await _post(app, "report.pdf", b"pdf-data", title="自定义标题")

    assert response.status_code == 200
    assert response.json()["data"]["title"] == "自定义标题"
    assert articles.insert_one.await_args.args[0]["title"] == "自定义标题"


@pytest.mark.asyncio
async def test_upload_requires_authentication() -> None:
    app, articles, _activities = _make_app(authenticated=False)

    response = await _post(app, "article.txt", ("正文" * 30).encode())

    assert response.status_code == 401
    assert response.json()["error"]["code"] == "NOT_AUTHENTICATED"
    articles.insert_one.assert_not_awaited()


@pytest.mark.asyncio
async def test_upload_rejects_unsupported_extension() -> None:
    app, articles, _activities = _make_app()

    response = await _post(app, "payload.exe", b"not-allowed")

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "UNSUPPORTED_FILE_TYPE"
    articles.insert_one.assert_not_awaited()


@pytest.mark.asyncio
async def test_upload_rejects_oversized_file() -> None:
    app, articles, _activities = _make_app()

    response = await _post(app, "large.txt", b"x" * (MAX_FILE_SIZE + 1))

    assert response.status_code == 413
    assert response.json()["error"]["code"] == "FILE_TOO_LARGE"
    articles.insert_one.assert_not_awaited()


@pytest.mark.asyncio
async def test_upload_rejects_empty_or_too_short_content() -> None:
    app, articles, _activities = _make_app()

    response = await _post(app, "short.md", b"# short")

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "EMPTY_CONTENT"
    articles.insert_one.assert_not_awaited()


@pytest.mark.asyncio
async def test_upload_maps_parser_failure_to_parse_failed() -> None:
    app, articles, _activities = _make_app()

    response = await _post(app, "broken.pdf", b"not-a-pdf")

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "PARSE_FAILED"
    articles.insert_one.assert_not_awaited()


@pytest.mark.asyncio
async def test_upload_rejects_duplicate_article() -> None:
    app, articles, activities = _make_app(duplicate=True)

    response = await _post(app, "article.txt", ("重复内容。" * 12).encode())

    assert response.status_code == 409
    assert response.json()["error"]["code"] == "DUPLICATE_ARTICLE"
    articles.insert_one.assert_not_awaited()
    activities.insert_one.assert_not_awaited()
